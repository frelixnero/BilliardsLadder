from pathlib import Path
from docx import Document

src = Path(r"c:/Users/USER/.gemini/antigravity/scratch/BilliardsLadder/docs/Production_Testing_Handoff_Report_May8_2026.md")
out = src.with_suffix('.docx')

lines = src.read_text(encoding='utf-8').splitlines()
doc = Document()

for line in lines:
    stripped = line.strip()

    if not stripped:
        doc.add_paragraph("")
        continue

    if stripped.startswith('### '):
        doc.add_heading(stripped[4:].strip(), level=3)
        continue
    if stripped.startswith('## '):
        doc.add_heading(stripped[3:].strip(), level=2)
        continue
    if stripped.startswith('# '):
        doc.add_heading(stripped[2:].strip(), level=1)
        continue

    if stripped.startswith('- '):
        doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        continue

    if stripped[:2].isdigit() and stripped[2:4] == '. ':
        doc.add_paragraph(stripped[4:].strip(), style='List Number')
        continue
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] == '. ':
        doc.add_paragraph(stripped[3:].strip(), style='List Number')
        continue

    if stripped == '---':
        doc.add_paragraph('----------------------------------------')
        continue

    doc.add_paragraph(stripped)

doc.save(out)
print(str(out))
