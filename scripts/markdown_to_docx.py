from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


INLINE_BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")


def add_inline_markdown(paragraph, text: str) -> None:
    parts = INLINE_BOLD_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def convert_markdown_to_docx(source: Path, target: Path) -> None:
    document = Document()
    lines = source.read_text(encoding="utf-8").splitlines()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped in {"---", "***", "___"}:
            document.add_paragraph("")
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[level:].strip()
            document.add_heading(heading_text, level=min(level, 4))
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:].strip())
            continue

        if re.match(r"^\d+\.\s+", stripped):
            item_text = re.sub(r"^\d+\.\s+", "", stripped)
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, item_text)
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, stripped)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python markdown_to_docx.py <source.md> <target.docx>")
        return 1

    source = Path(sys.argv[1]).resolve()
    target = Path(sys.argv[2]).resolve()

    if not source.exists():
        print(f"Source file not found: {source}")
        return 1

    convert_markdown_to_docx(source, target)
    print(f"Created: {target}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())